from PyPDF2 import PdfReader
import pdfplumber
from langchain_core.documents import Document
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from models import load_embeddings
import glob
import docx 
import os
import zipfile 
import xml.etree.ElementTree as ET 
import pandas as pd
import re

def load_all_documents_to_list(directory_path):
    # 폴더가 없으면 빈 리스트 반환 (에러 방지)
    if not os.path.exists(directory_path):
        print(f"경로 없음: {directory_path}")
        return []

    all_documents = glob.glob(os.path.join(directory_path, "*"))
    documents = []

    for file_path in all_documents:
        try:
            if os.path.isdir(file_path):
                continue

            elif file_path.endswith(".pdf"):  #pdf로 끝나는 경우
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        text = page.extract_text() or ""
                        tables = page.extract_tables()
                        table_text = ""
                        if tables:
                            for table in tables:
                                table_text += "\n\n--- TABLE START ---\n"
                                for row in table:
                                    # None 값 처리
                                    safe_row = [str(item) if item is not None else "" for item in row]
                                    table_text += " | ".join(safe_row) + "\n"
                                table_text += "--- TABLE END ---\n\n"
                        
                        full_content = text + table_text
                        documents.append(Document(page_content=full_content, metadata={"source": file_path, "page": page_num}))

            elif file_path.endswith(".docx"): #docx 파일 
                doc = docx.Document(file_path)
                full_text = "\n".join([para.text for para in doc.paragraphs])
                documents.append(Document(page_content=full_text, metadata={"source": file_path}))

            elif file_path.endswith(".hwpx"): #hwpx 파일
                full_text = ""
                with zipfile.ZipFile(file_path, 'r') as z:
                    section_xml_files = [f for f in z.namelist() if f.startswith('Contents/section') and f.endswith('.xml')]
                    for section_file in sorted(section_xml_files):
                        xml_content = z.read(section_file)
                        root = ET.fromstring(xml_content)
                        for text_element in root.iter('{http://www.hancom.co.kr/hwpml/2011/paragraph}t'):
                            if text_element.text:
                                full_text += text_element.text + "\n"
                documents.append(Document(page_content=full_text, metadata={"source": file_path}))

            elif file_path.endswith(".txt"): #txt 파일
                with open(file_path, 'r', encoding='utf-8') as f:
                    raw_text = f.read()
                    pattern = r'^.*사진.*$\n?'
                    full_text = re.sub(pattern, '', raw_text, flags=re.MULTILINE)
                documents.append(Document(page_content=full_text, metadata={"source": os.path.basename(file_path)}))
            
            elif file_path.endswith(".xlsx"): 
                # 헤더 0 (첫줄), A~E열 사용
                df = pd.read_excel(file_path, sheet_name=0, header=0, usecols="A:E")
                for index, row in df.iterrows():
                    row_content = ", ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                    doc = Document(
                        page_content=row_content,
                        metadata={"source": os.path.basename(file_path), "row": index + 2}
                    )
                    documents.append(doc)

        except Exception as e:
            print(f"Error processing {file_path}: {e}")
            continue

    return documents 

def build_vectorstores():
    """문서 분할 + 벡터DB 생성"""
    
    # 데이터 로드
    law_docs = load_all_documents_to_list("Dataset/법령제도데이터")
    manual_docs = load_all_documents_to_list("Dataset/매뉴얼데이터")
    basic_docs = load_all_documents_to_list("Dataset/기본데이터")
    past_docs = load_all_documents_to_list("Dataset/이력데이터")

    # 데이터가 없으면 임시(Dummy) 문서 생성
    if not law_docs:
        print(" [법령] 데이터 없음 -> 임시 데이터 생성")
        law_docs = [Document(page_content="법령 정보가 아직 없습니다.", metadata={"source": "dummy_law"})]
        
    if not manual_docs:
        print("[매뉴얼] 데이터 없음 -> 임시 데이터 생성")
        manual_docs = [Document(page_content="매뉴얼 정보가 아직 없습니다.", metadata={"source": "dummy_manual"})]
        
    if not basic_docs:
        print("[기본] 데이터 없음 -> 임시 데이터 생성")
        basic_docs = [Document(page_content="기본 정보가 아직 없습니다.", metadata={"source": "dummy_basic"})]
        
    if not past_docs:
        print("[이력] 데이터 없음 -> 임시 데이터 생성")
        past_docs = [Document(page_content="이력 정보가 아직 없습니다.", metadata={"source": "dummy_past"})]

    # 스플리터 설정
    splitter = CharacterTextSplitter(chunk_size=200, chunk_overlap=0)

    law_splits = splitter.split_documents(law_docs)
    manual_splits = splitter.split_documents(manual_docs)
    basic_splits = splitter.split_documents(basic_docs)
    past_splits = splitter.split_documents(past_docs)

    # 임베딩 로드
    embeddings = load_embeddings()

    # 벡터 DB 생성
    vectordb_law = FAISS.from_documents(law_splits, embeddings)
    vectordb_manual = FAISS.from_documents(manual_splits, embeddings)
    vectordb_basic = FAISS.from_documents(basic_splits, embeddings)
    vectordb_past = FAISS.from_documents(past_splits, embeddings)

    return vectordb_law, vectordb_manual, vectordb_basic, vectordb_past